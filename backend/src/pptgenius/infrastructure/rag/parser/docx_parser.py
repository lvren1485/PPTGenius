""".docx parser via python-docx."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from .base import ParsedDocument, _normalise_text


def parse_docx(file_path: str | Path) -> ParsedDocument:
    """Extract text from a .docx file."""
    path = Path(file_path)
    doc = Document(str(path))

    paragraphs: list[str] = []
    for para in doc.paragraphs:
        paragraphs.append(para.text)

    # Also extract tables as Markdown
    for table in doc.tables:
        rows: list[str] = []
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
            if i == 0:
                rows.append("| " + " | ".join(["---"] * len(row.cells)) + " |")
        paragraphs.append("")
        paragraphs.extend(rows)
        paragraphs.append("")

    raw = "\n".join(paragraphs)
    text = _normalise_text(raw)
    return ParsedDocument(
        file_path=str(path),
        file_type="docx",
        text=text,
        metadata={
            "filename": path.name,
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
        },
    )
